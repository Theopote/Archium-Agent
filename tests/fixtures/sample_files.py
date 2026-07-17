"""Generate sample files for parser and ingestion tests."""

from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook
from PIL import Image


def create_sample_pdf(path: Path, *, text: str = "Site traffic analysis summary.") -> Path:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    document.save(path)
    document.close()
    return path


def create_sample_docx(path: Path, *, heading: str = "现状分析", body: str = "项目背景说明。") -> Path:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(body)
    document.save(path)
    return path


def create_sample_xlsx(path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "指标"
    sheet.append(["指标", "数值"])
    sheet.append(["用地面积", "12000"])
    workbook.save(path)
    return path


def create_sample_image(path: Path) -> Path:
    image = Image.new("RGB", (800, 600), color=(120, 140, 160))
    image.save(path, format="JPEG")
    return path
